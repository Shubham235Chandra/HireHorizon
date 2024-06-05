import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
import json
from docx import Document
from dotenv import load_dotenv
import io

load_dotenv()  # Load all the environment variables

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def get_gemini_response(resume_text, jd):
    # Format the input prompt with the actual data
    formatted_prompt = input_prompt.format(text=resume_text, jd=jd)
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(formatted_prompt)
    return response.text

def input_pdf_text(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:  # Iterate over each page in the PDF
        text += page.extract_text() if page.extract_text() else ""  # Append the extracted text
    return text

def input_docx_text(uploaded_file):
    doc = Document(uploaded_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def update_resume_docx(modified_text):
    modified_doc = Document()

    # Add the modified text to the new document
    for line in modified_text.split("\n"):
        modified_doc.add_paragraph(line)

    return modified_doc

def clean_json_response(response_text):
    try:
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        cleaned_text = response_text[json_start:json_end]
        return json.loads(cleaned_text)
    except Exception as e:
        st.error(f"Failed to clean JSON response: {e}")
        return None

# Prompt
input_prompt = """
Act as an advanced ATS (Application Tracking System) with deep expertise in the tech field, including software engineering, data science, data analytics, and big data engineering. 
Your primary task is to evaluate the resume in comparison with the provided job description. 
Given the highly competitive job market, your analysis should offer detailed assistance to improve the resume.

1. Resume Analysis: Provide a score reflecting the percentage match between the resume and the job description.
2. Keyword Identification: Identify critical missing keywords from the original resume when compared to the job description.
3. Profile Summary: Offer a concise profile summary based on an analysis that could enhance the applicant's alignment with the job, including the incorporation of previously missing keywords.
4. Resume Modification**: Provide a modified version of the resume, maintaining the original format but incorporating all identified missing keywords in the most effective places to improve ATS compatibility. This modification should adhere to ATS best practices and should not alter unnecessary details or the overall format of the resume.

Input parameters:
- Resume Text: {text}
- Job Description: {jd}


Response format:
{{
  "JD Match": "%",
  "Missing Keywords": [],
  "Profile Summary": "",
  "Modified Resume": ""
}}
"""

## Streamlit App
st.title("Hire Horizon")
st.text("Improve Your Resume ATS")
jd = st.text_area("Paste the Job Description")
uploaded_file = st.file_uploader("Upload Your Resume", type=["pdf", "docx"], help="Please upload the PDF or DOCX of your resume")
st.text("Please upload the PDF or DOCX of your Resume")

submit = st.button("Submit")

if submit:
    if uploaded_file is not None and jd:
        if uploaded_file.type == "application/pdf":
            resume_text = input_pdf_text(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = input_docx_text(uploaded_file)
        else:
            st.error("Unsupported file type. Please upload a PDF or DOCX file.")
            resume_text = None

        if resume_text:
            response = get_gemini_response(resume_text, jd)
            response_data = clean_json_response(response)
            if response_data:
                # Display each component separately
                st.subheader("Profile Summary")
                st.write(response_data['Profile Summary'] if response_data['Profile Summary'] else "No additional profile summary provided.")

                st.subheader("Job Description Match")
                st.write(f"{response_data['JD Match']} Match")
                
                st.subheader("Missing Keywords")
                if response_data['Missing Keywords']:
                    st.write(response_data['Missing Keywords'])
                else:
                    st.write("Resume Seems Perfect")

                # Display modified resume
                st.subheader("Modified Resume")
                st.text_area("Modified Resume Content", response_data['Modified Resume'], height=400)

                # Generate updated resume file
                if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    updated_resume = update_resume_docx(response_data['Modified Resume'])
                    
                    # Save updated docx to a BytesIO object
                    buffer = io.BytesIO()
                    updated_resume.save(buffer)
                    buffer.seek(0)
                    
                    # Display the updated resume
                    st.subheader("Updated Resume")
                    updated_resume_text = input_docx_text(buffer)
                    st.text_area("Updated Resume Content", updated_resume_text, height=400)
            else:
                st.error("Failed to parse the response from the AI model.")

